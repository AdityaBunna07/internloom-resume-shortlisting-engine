"""Fault-tolerant PDF resume parsing for InternLoom."""

from __future__ import annotations

import io
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from skills_vocab import SKILL_ALIASES


SECTION_HEADERS = {
    "projects": ("projects", "project experience", "academic projects", "personal projects"),
    "experience": ("experience", "work experience", "employment", "internship", "internships"),
    "certifications": ("certifications", "certificates", "courses", "achievements"),
}

SUPPORTED_RESUME_EXTENSIONS = {".pdf", ".doc", ".docx"}


def _empty_resume(path: Path) -> dict[str, Any]:
    return {
        "filename": path.name,
        "full_name": None,
        "email": None,
        "phone": None,
        "college": None,
        "degree_branch": None,
        "graduation_year": None,
        "cgpa": None,
        "cgpa_raw": None,
        "cgpa_scale_assumed": False,
        "skills": [],
        "projects": [],
        "experience": [],
        "certifications": [],
        "parse_status": "Failed",
        "parse_reason": "No usable text could be extracted.",
    }


def _clean(value: str | None) -> str | None:
    if not value:
        return None
    value = re.sub(r"\s+", " ", value).strip(" -|:\t")
    return value or None


def _extract_pdfplumber_text(path: Path) -> str:
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception:
        return ""


def _column_ordered_text(blocks: list[tuple], page_width: float) -> tuple[str, bool]:
    usable = [block for block in blocks if len(block) >= 5 and str(block[4]).strip()]
    if len(usable) < 4:
        return "", False
    positions = sorted(set(round(float(block[0]), 1) for block in usable))
    if len(positions) < 2:
        return "", False
    gaps = [(positions[index + 1] - positions[index], index) for index in range(len(positions) - 1)]
    largest_gap, gap_index = max(gaps)
    split = (positions[gap_index] + positions[gap_index + 1]) / 2
    left = [block for block in usable if float(block[0]) < split]
    right = [block for block in usable if float(block[0]) >= split]
    confident = (
        largest_gap >= page_width * 0.16
        and len(left) >= 2
        and len(right) >= 2
        and abs(sum(float(block[0]) for block in left) / len(left) - sum(float(block[0]) for block in right) / len(right)) >= page_width * 0.2
    )
    if not confident:
        return "", False
    sort_key = lambda block: (round(float(block[1]), 1), float(block[0]))
    ordered = sorted(left, key=sort_key) + sorted(right, key=sort_key)
    return "\n".join(str(block[4]).strip() for block in ordered), True


def _extract_fitz_text(path: Path) -> tuple[str, bool]:
    try:
        import fitz

        document = fitz.open(path)
        page_texts: list[str] = []
        column_detected = False
        for page in document:
            text, is_column = _column_ordered_text(page.get_text("blocks"), float(page.rect.width))
            page_texts.append(text if is_column else page.get_text("text"))
            column_detected = column_detected or is_column
        document.close()
        return "\n".join(page_texts), column_detected
    except Exception:
        return "", False


def _ocr_text(path: Path) -> str:
    try:
        import fitz
        import pytesseract
        from PIL import Image

        document = fitz.open(path)
        pages = []
        for page in document:
            pixmap = page.get_pixmap(dpi=300, alpha=False)
            pages.append(pytesseract.image_to_string(Image.open(io.BytesIO(pixmap.tobytes("png")))))
        document.close()
        return "\n".join(pages)
    except Exception:
        return ""


def _extract_pdf_text(path: Path) -> tuple[str, str | None]:
    plumber_text = _extract_pdfplumber_text(path)
    fitz_text, used_columns = _extract_fitz_text(path)
    text = fitz_text if used_columns and len(fitz_text.strip()) >= 50 else plumber_text or fitz_text
    if len(text.strip()) >= 50:
        return text, None
    ocr_text = _ocr_text(path)
    if len(ocr_text.strip()) >= 20:
        return ocr_text, "OCR fallback used because the PDF had little embedded text."
    return "", "No usable embedded text or OCR output; PDF may be corrupt, protected, or image-only."


def _extract_docx_text(path: Path) -> tuple[str, str | None]:
    try:
        from docx import Document

        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(part for part in parts if part.strip())
        if len(text.strip()) >= 20:
            return text, None
        return "", "The DOCX file contains too little readable text."
    except Exception as error:
        return "", f"Could not read DOCX file: {type(error).__name__}."


def _extract_legacy_doc_text(path: Path) -> tuple[str, str | None]:
    for command in (("antiword", str(path)), ("catdoc", str(path))):
        executable = shutil.which(command[0])
        if not executable:
            continue
        try:
            process = subprocess.run((executable, command[1]), capture_output=True, text=True, timeout=30, check=False)
            if process.returncode == 0 and len(process.stdout.strip()) >= 20:
                return process.stdout, None
        except Exception:
            continue
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            with tempfile.TemporaryDirectory() as temporary_directory:
                process = subprocess.run((soffice, "--headless", "--convert-to", "txt:Text", "--outdir", temporary_directory, str(path)), capture_output=True, text=True, timeout=45, check=False)
                converted = Path(temporary_directory) / f"{path.stem}.txt"
                if process.returncode == 0 and converted.exists():
                    text = converted.read_text(encoding="utf-8", errors="ignore")
                    if len(text.strip()) >= 20:
                        return text, None
        except Exception:
            pass
    try:
        import pythoncom
        import win32com.client

        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(path.resolve()), ReadOnly=True)
        text = document.Content.Text
        document.Close(False)
        word.Quit()
        if len(text.strip()) >= 20:
            return text, None
    except Exception:
        pass
    return "", "Legacy .doc extraction needs Microsoft Word, LibreOffice, antiword, or catdoc on this server."


def _extract_text(path: Path) -> tuple[str, str | None]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)
    if suffix == ".doc":
        return _extract_legacy_doc_text(path)
    return "", f"Unsupported resume format: {suffix or 'no extension'}."


def _find_name(text: str) -> str | None:
    section_labels = {
        "professional summary",
        "summary",
        "skills",
        "technical skills",
        "work experience",
        "experience",
        "education",
        "projects",
        "certifications",
        "career objective",
        "profile",
    }
    context_terms = {
        "language",
        "skills",
        "college",
        "university",
        "cgpa",
        "gpa",
        "experience",
        "project",
        "developer",
        "engineer",
        "performance",
        "technology",
        "education",
    }
    for line in text.splitlines()[:3]:
        candidate = _clean(line)
        if not candidate or len(candidate) > 70 or "@" in candidate or re.search(r"\d{4,}", candidate):
            continue
        if candidate.lower() in section_labels:
            continue
        words = re.findall(r"[A-Za-z][A-Za-z.'-]*", candidate)
        if (
            2 <= len(words) <= 5
            and all(word.lower() not in {"resume", "curriculum", "vitae", "profile"} for word in words)
            and not any(word.lower() in context_terms for word in words)
            and all(word[0].isupper() for word in words)
        ):
            return " ".join(words)
    return None


def _find_first(text: str, patterns: tuple[str, ...]) -> str | None:
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if match:
            return _clean(match.group(1))
    return None


def _parse_cgpa(text: str) -> tuple[float | None, str | None, bool]:
    patterns = (
        r"\b(?:cgpa|gpa)\s*[:=-]?\s*(\d{1,2}(?:\.\d{1,2})?)\s*(/\s*(?:10|4)|%)?",
        r"\b(\d{1,2}(?:\.\d{1,2})?)\s*(/\s*(?:10|4)|%)\s*(?:cgpa|gpa)?",
    )
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if not match:
            continue
        value = float(match.group(1))
        marker = (match.group(2) or "").replace(" ", "").lower()
        raw = match.group(0).strip()
        if marker == "%":
            return round(value / 9.5, 2), raw, False
        if marker == "/4":
            return round(value * 2.5, 2), raw, False
        if marker == "/10":
            return round(value, 2), raw, False
        assumed = value <= 4.0
        return round(value * 2.5 if assumed else value, 2), raw, True
    return None, None, False


def _extract_skills(text: str) -> list[str]:
    matches = []
    lowered = text.lower()
    for canonical, aliases in SKILL_ALIASES.items():
        for alias in aliases:
            pattern = r"(?<![a-z0-9+#.])" + re.escape(alias.lower()) + r"(?![a-z0-9+#.])"
            if re.search(pattern, lowered):
                matches.append(canonical)
                break
    return sorted(set(matches))


def _section_lines(text: str, section: str) -> list[str]:
    lines = [_clean(line) for line in text.splitlines()]
    lines = [line for line in lines if line]
    headers = SECTION_HEADERS[section]
    start = next((index for index, line in enumerate(lines) if line.lower().rstrip(":") in headers), None)
    if start is None:
        return []
    result = []
    for line in lines[start + 1 : start + 22]:
        if line.lower().rstrip(":") in sum((list(value) for value in SECTION_HEADERS.values()), []):
            break
        result.append(line)
    return result


def _extract_projects(text: str) -> list[dict[str, str]]:
    lines = _section_lines(text, "projects")
    projects = []
    for line in lines:
        item = re.sub(r"^[•*-]\s*", "", line)
        if len(item) < 4:
            continue
        pieces = re.split(r"\s*(?:[:|–—-])\s*", item, maxsplit=1)
        title = _clean(pieces[0])
        if title and len(title) <= 90:
            projects.append({"title": title, "description": _clean(pieces[1]) if len(pieces) > 1 else ""})
        if len(projects) == 8:
            break
    return projects


def _extract_experience(text: str) -> list[dict[str, str | None]]:
    lines = _section_lines(text, "experience")
    results = []
    for line in lines:
        item = re.sub(r"^[•*-]\s*", "", line)
        if len(item) < 4:
            continue
        parts = [part.strip() for part in re.split(r"\s*[|–—]\s*", item) if part.strip()]
        role = parts[0] if parts else None
        company = parts[1] if len(parts) > 1 else None
        duration_match = re.search(r"\b(?:\d+\s*(?:month|mos|year|yr)s?|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s*20\d{2}\s*(?:-|to)\s*(?:present|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s*20\d{2}))", item, re.I)
        results.append({"company": company, "role": role, "duration": duration_match.group(0) if duration_match else None})
        if len(results) == 8:
            break
    return results


def _extract_certifications(text: str) -> list[str]:
    return [re.sub(r"^[•*-]\s*", "", line) for line in _section_lines(text, "certifications")[:10]]


def parse_resume(resume_path: str | Path) -> dict[str, Any]:
    path = Path(resume_path)
    resume = _empty_resume(path)
    try:
        text, extraction_note = _extract_text(path)
        if not text:
            resume["parse_reason"] = extraction_note or resume["parse_reason"]
            return resume
        email_match = re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", text, re.I)
        phone_match = re.search(r"(?<!\d)(?:\+?91[\s-]?)?(?:\(?\d{3,5}\)?[\s-]?)?\d{5}[\s-]?\d{5}(?!\d)", text)
        cgpa, cgpa_raw, assumed = _parse_cgpa(text)
        resume.update(
            full_name=_find_name(text),
            email=email_match.group(0) if email_match else None,
            phone=phone_match.group(0) if phone_match else None,
            college=_find_first(text, (r"(?:college|university|institute)\s*(?:of|for)?\s*([A-Za-z][^\n,]{2,90})",)),
            degree_branch=_find_first(text, (r"\b((?:b\.?tech|b\.?e\.?|bachelor(?:'s)?|m\.?tech|m\.?c\.?a)[^\n,;]{0,80})",)),
            graduation_year=_find_first(text, (r"(?:graduat(?:ion|ing)|expected|batch|pass(?:ing)?\s*year)\s*[:=-]?\s*((?:19|20)\d{2})",)),
            cgpa=cgpa,
            cgpa_raw=cgpa_raw,
            cgpa_scale_assumed=assumed,
            skills=_extract_skills(text),
            projects=_extract_projects(text),
            experience=_extract_experience(text),
            certifications=_extract_certifications(text),
        )
        key_fields = ("full_name", "email", "degree_branch", "graduation_year", "cgpa")
        missing = [field.replace("_", " ") for field in key_fields if not resume[field]]
        if not missing and resume["skills"]:
            resume["parse_status"] = "Clean"
            resume["parse_reason"] = None
        else:
            details = []
            if missing:
                details.append("Missing " + ", ".join(missing) + ".")
            if not resume["skills"]:
                details.append("No vocabulary skills detected.")
            if extraction_note:
                details.append(extraction_note)
            resume["parse_status"] = "Partial"
            resume["parse_reason"] = " ".join(details)
        return resume
    except Exception as error:
        resume["parse_reason"] = f"Unexpected parser error: {type(error).__name__}."
        return resume


def parse_resumes(resume_folder: str | Path) -> list[dict[str, Any]]:
    folder = Path(resume_folder)
    if not folder.exists() or not folder.is_dir():
        raise ValueError(f"Resume folder does not exist or is not a directory: {folder}")
    return [parse_resume(path) for path in sorted(folder.iterdir()) if path.is_file() and path.suffix.lower() in SUPPORTED_RESUME_EXTENSIONS]
